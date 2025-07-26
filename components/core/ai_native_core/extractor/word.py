import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import List, Union, Dict, Any
from urllib.parse import urlparse
from xml.etree import ElementTree

import requests
from docx import Document as DocxDocument
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from langchain_core.messages import SystemMessage, HumanMessage

from ai_native_core.extensions.ext_redis import redis_client
from ai_native_core.extensions.ext_storage import storage
from ai_native_core.model import last_model


class VisionPictureConfig:
    """Configuration for vision picture understanding."""

    def __init__(
            self,
            model: str,
            model_provider: str,
            base_url: str,
            api_key: str,
            max_tokens: int = 4096,
            temperature: float = 0,
            system_prompt: str = "你是一个图片内容提取专家，擅长提取图片中的内容，理解图片中的内容。保留图片中的所有细节和信息，尽可能全面地描述图片中的内容。用markdown的格式输出图片内容，包含图片的描述和相关信息。",
            human_prompt: str = "用你的才能提取下面图片中的内容",
    ):
        self.model = model
        self.model_provider = model_provider
        self.base_url = base_url
        self.api_key = api_key
        self.max_tokens = max_tokens
        self.temperature = temperature
        self.system_prompt = system_prompt
        self.human_prompt = human_prompt


class DocxLoader(BaseLoader):

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def __init__(
            self,
            file_path: Union[str, Path],
            tenant_id: str,
            document_id: str,
            is_vision_picture_understanding: bool = False,
            vision_model_config: VisionPictureConfig = None,
    ):
        """Initialize with file path."""
        self.file_path = str(file_path)
        self.tenant_id = tenant_id
        self.document_id = document_id
        self.original_file_path = self.file_path
        self.is_vision_picture_understanding = is_vision_picture_understanding
        # to dict
        self.vision_model_config = vision_model_config
        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # If the file is a web path, download it to a temporary file, and use that
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            if r.status_code != 200:
                raise ValueError(
                    "Check the url of your file; returned status code %s"
                    % r.status_code
                )

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile()
            self.temp_file.write(r.content)
            self.file_path = self.temp_file.name
        elif not os.path.isfile(self.file_path):
            raise ValueError("File path %s is not a valid file or url" % self.file_path)

    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def load(self) -> List[Document]:
        """Load given path as single page."""
        content = self.parse_docx(self.file_path)
        return [
            Document(
                page_content=content,
                metadata={"source": self.file_path},
            )
        ]

    def _extract_images_from_docx(self, doc) -> Dict[Any, str]:
        """Extract images from docx and return a mapping of image parts to markdown links."""
        image_count = 0
        image_map = {}

        # Redis key for storing image UUIDs of this document
        redis_key = f"image_files:{self.tenant_id}:{self.document_id}"

        # 获取当前文档已存在的图片UUID集合
        existing_image_uuids = set()
        try:
            existing_data = redis_client.smembers(redis_key)
            if existing_data:
                existing_image_uuids = {uuid_bytes.decode('utf-8') for uuid_bytes in existing_data}
        except Exception as e:
            print(f"Failed to get existing images from Redis: {e}")

        # 当前这次处理的图片UUID集合
        current_image_uuids = set()

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                # 图像是外部的就进行下载
                if rel.is_external:
                    # TODO: 处理外部的图片
                    pass
                else:
                    image_ext = rel.target_ref.split(".")[-1]
                    if image_ext is None:
                        continue

                    # 生成图片内容的hash值作为去重依据
                    image_data = rel.target_part.blob

                    # 检查是否已存在相同内容的图片
                    existing_uuid = None
                    for uuid_str in existing_image_uuids:
                        # 构造可能的文件key来检查是否存在
                        try:
                            # 简化处理：如果redis中存在这个uuid，我们认为图片已存在
                            existing_uuid = uuid_str
                            break
                        except:
                            continue

                    if existing_uuid:
                        # 使用已存在的图片UUID
                        file_uuid = existing_uuid
                    else:
                        # 生成新的UUID
                        file_uuid = str(uuid.uuid4())
                        file_key = f"image_files/{self.tenant_id}/{self.document_id}/{file_uuid}.{image_ext}"

                        # 增量存储到对象存储中
                        try:
                            storage.save(file_key, image_data)
                            print(f"Saved new image: {file_key}")
                        except Exception as e:
                            print(f"Failed to save image {file_key}: {e}")
                            continue
                    # 将当前使用的图片UUID加入到当前集合
                    current_image_uuids.add(file_uuid)

                    # 生成 markdown 格式的图片链接
                    image_url = (
                        f"{os.environ.get('STORAGE_STATIC_RESOURCE_PREFIX_URL')}/"
                        f"image_files/{self.tenant_id}/{self.document_id}/{file_uuid}.{image_ext}"
                    )
                    image_map[rel.target_part] = f"![image]({image_url})\n\n"

                    if self.is_vision_picture_understanding:
                        # base64_image = base64.b64encode(image_data).decode('utf-8')
                        system_message = SystemMessage(
                            content=self.vision_model_config.system_prompt
                        )
                        message = HumanMessage(
                            content=[
                                {"type": "text", "text": self.vision_model_config.human_prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {"url": image_url},
                                },
                            ]
                        )
                        picture_content = last_model.invoke(
                            [
                                system_message,
                                message
                            ],
                            config={
                                "configurable": {
                                    "last_base_url": self.vision_model_config.base_url,
                                    "last_model": self.vision_model_config.model,
                                    "last_model_provider": self.vision_model_config.model_provider,
                                    "last_temperature": self.vision_model_config.temperature,
                                    "last_max_tokens": self.vision_model_config.max_tokens,
                                    "last_api_key": self.vision_model_config.api_key
                                }
                            }
                        )
                        # 将图片内容添加到markdown链接中
                        image_map[rel.target_part] += f"\n\n上面图片描述的内容为:{picture_content.content}\n\n"

        # 更新Redis中的图片UUID集合
        try:
            # 删除旧的集合
            redis_client.delete(redis_key)
            # 添加当前使用的图片UUID
            if current_image_uuids:
                redis_client.sadd(redis_key, *current_image_uuids)
                # 设置过期时间(30天)，避免Redis中数据过多
                redis_client.expire(redis_key, 30 * 24 * 3600)
        except Exception as e:
            print(f"Failed to update Redis: {e}")

        # 清理不再使用的图片
        try:
            obsolete_uuids = existing_image_uuids - current_image_uuids
            if obsolete_uuids:
                self._cleanup_obsolete_images(obsolete_uuids)
        except Exception as e:
            print(f"Failed to cleanup obsolete images: {e}")

        return image_map

    def _cleanup_obsolete_images(self, obsolete_uuids: set) -> None:
        """清理不再使用的图片"""
        for uuid_str in obsolete_uuids:
            # 可能的图片扩展名
            possible_extensions = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']
            for ext in possible_extensions:
                file_key = f"image_files/{self.tenant_id}/{self.document_id}/{uuid_str}.{ext}"
                try:
                    # 删除对象存储中的文件
                    storage.delete(file_key)
                    print(f"Deleted obsolete image: {file_key}")
                except Exception as e:
                    # 如果文件不存在或删除失败，继续处理下一个
                    continue

    @staticmethod
    def _get_heading_level(paragraph) -> int:
        """Get the heading level of a paragraph (1-6 for h1-h6, 0 for non-heading)."""
        style = getattr(paragraph, "style", None)
        style_name = getattr(style, "name", None)
        if not style_name:
            return 0
        style_name = style_name.lower()

        # 检查是否为标题样式
        if 'heading' in style_name:
            try:
                # 提取标题级别
                match = re.search(r'heading\s*(\d+)', style_name)
                if match:
                    level = int(match.group(1))
                    return min(level, 6)  # Markdown 最多支持 6 级标题
            except (AttributeError, ValueError):
                pass

        # 检查其他可能的标题标识
        if any(keyword in style_name for keyword in ['title', 'subtitle', 'header']):
            return 1

        return 0

    @staticmethod
    def _format_text_run(run) -> str:
        """Format a text run with markdown formatting."""
        text = run.text
        if not text.strip():
            return text

        # 处理粗体
        if run.bold:
            text = f"**{text}**"

        # 处理斜体
        if run.italic:
            text = f"*{text}*"

        # 处理下划线（用粗体表示）
        if run.underline:
            text = f"__{text}__"

        # 处理删除线
        if hasattr(run, 'strike') and run.strike:
            text = f"~~{text}~~"

        return text

    @staticmethod
    def _is_list_paragraph(paragraph) -> tuple:
        """Check if paragraph is a list item and return (is_list, is_ordered, level)."""
        # 检查段落的编号格式
        if paragraph._element.pPr is not None:
            numPr = paragraph._element.pPr.find(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                # 获取列表级别
                ilvl = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                level = int(ilvl.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if ilvl is not None else 0

                # 获取编号ID来判断是否为有序列表
                numId = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                if numId is not None:
                    # 这里简化处理，通常编号列表为有序列表
                    return True, True, level

        # 检查样式名称中是否包含列表信息
        style = getattr(paragraph, "style", None)
        style_name = getattr(style, "name", None)
        if style_name:
            style_name = style_name.lower()
            if 'list' in style_name:
                if 'bullet' in style_name:
                    return True, False, 0
                elif 'number' in style_name:
                    return True, True, 0

        return False, False, 0

    def _table_to_markdown(self, table, image_map) -> str:
        """Convert a table to markdown format."""
        markdown = []

        # 计算表格的总列数
        total_cols = max(len(row.cells) for row in table.rows)

        # 处理表头
        if table.rows:
            header_row = table.rows[0]
            headers = self._parse_table_row(header_row, image_map, total_cols)
            markdown.append("| " + " | ".join(headers) + " |")
            markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

            # 处理数据行
            for row in table.rows[1:]:
                row_cells = self._parse_table_row(row, image_map, total_cols)
                markdown.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(markdown)

    def _parse_table_row(self, row, image_map, total_cols) -> List[str]:
        """Parse a table row and return cell contents."""
        row_cells = [""] * total_cols
        col_index = 0

        for cell in row.cells:
            # 确保列索引不超出范围
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1

            if col_index >= total_cols:
                break

            cell_content = self._parse_table_cell(cell, image_map).strip()
            cell_colspan = getattr(cell, 'grid_span', 1) or 1

            # 处理合并单元格
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""

            col_index += cell_colspan

        return row_cells

    def _parse_table_cell(self, cell, image_map) -> str:
        """Parse a table cell and return formatted content."""
        cell_content = []

        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_paragraph_content(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)

        # 去重并连接内容
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_paragraph_content(self, paragraph, image_map) -> str:
        """Parse paragraph content with formatting and images."""
        paragraph_content = []

        for run in paragraph.runs:
            # 处理图片
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed_id:
                        try:
                            image_part = run.part.related_parts.get(embed_id)
                            if image_part in image_map:
                                paragraph_content.append(image_map[image_part])
                        except:
                            pass

            # 处理普通文本
            if run.text.strip():
                formatted_text = self._format_text_run(run)
                paragraph_content.append(formatted_text)

        return "".join(paragraph_content).strip()

    @staticmethod
    def _process_hyperlinks(doc) -> None:
        """Process hyperlinks in the document."""
        hyperlinks_url = None
        url_pattern = re.compile(r"https?://[^\s]+")

        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and hyperlinks_url:
                    result = f"[{run.text}]({hyperlinks_url})"
                    run.text = result
                    hyperlinks_url = None

                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ElementTree.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x is None:
                                continue
                            if x.tag.endswith("instrText"):
                                if x.text is None:
                                    continue
                                for i in url_pattern.findall(x.text):
                                    hyperlinks_url = str(i)
                    except Exception:
                        pass

    def parse_docx(self, docx_path: str) -> str:
        """Parse a docx file and return markdown content."""
        doc = DocxDocument(docx_path)

        content = []
        image_map = self._extract_images_from_docx(doc)

        # 处理超链接
        self._process_hyperlinks(doc)

        # 获取所有段落和表格，保持原有顺序
        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()

        # 遍历文档的所有元素
        prev_element_type = None  # 追踪前一个元素类型

        for element in doc.element.body:
            if not hasattr(element, "tag"):
                continue

            if isinstance(element.tag, str) and element.tag.endswith("p"):  # 段落
                if paragraphs:
                    para = paragraphs.pop(0)

                    # 检查是否为标题
                    heading_level = self._get_heading_level(para)
                    if heading_level > 0:
                        paragraph_text = self._parse_paragraph_content(para, image_map)
                        if paragraph_text.strip():
                            # 在标题前添加空行（除非是第一个元素）
                            if content and prev_element_type != "heading":
                                content.append("")
                            content.append(f"{'#' * heading_level} {paragraph_text}")
                            # 在标题后添加空行
                            content.append("")
                            prev_element_type = "heading"
                        continue

                    # 检查是否为列表
                    is_list, is_ordered, level = self._is_list_paragraph(para)
                    if is_list:
                        paragraph_text = self._parse_paragraph_content(para, image_map)
                        if paragraph_text.strip():
                            # 如果前一个元素不是列表，添加空行
                            if content and prev_element_type not in ["list", "heading"]:
                                content.append("")
                            indent = "  " * level
                            list_marker = "1." if is_ordered else "-"
                            content.append(f"{indent}{list_marker} {paragraph_text}")
                            prev_element_type = "list"
                        continue

                    # 普通段落
                    paragraph_text = self._parse_paragraph_content(para, image_map)
                    if paragraph_text.strip():
                        # 在普通段落前添加空行（除非前一个是标题或者是第一个元素）
                        if content and prev_element_type not in ["paragraph", "heading"]:
                            content.append("")
                        content.append(paragraph_text)
                        # 在普通段落后添加空行
                        content.append("")
                        prev_element_type = "paragraph"
                    else:
                        # 保持空行
                        content.append("")

            elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # 表格
                if tables:
                    table = tables.pop(0)
                    table_markdown = self._table_to_markdown(table, image_map)
                    # 在表格前添加空行
                    if content:
                        content.append("")
                    content.append(table_markdown)
                    # 在表格后添加空行
                    content.append("")
                    prev_element_type = "table"

        # 清理和格式化输出
        formatted_content = []
        prev_line_empty = False
        consecutive_empty_lines = 0

        for line in content:
            if line.strip() == "":
                consecutive_empty_lines += 1
                # 最多保留一个连续的空行
                if consecutive_empty_lines == 1:
                    formatted_content.append("")
                prev_line_empty = True
            else:
                formatted_content.append(line)
                prev_line_empty = False
                consecutive_empty_lines = 0

        # 移除开头和结尾的空行
        while formatted_content and formatted_content[0].strip() == "":
            formatted_content.pop(0)
        while formatted_content and formatted_content[-1].strip() == "":
            formatted_content.pop()

        return "\n".join(formatted_content)
